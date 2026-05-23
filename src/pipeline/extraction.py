"""PDF/DOCX text extraction — Step ① of the contract diff pipeline."""

import os
import pdfplumber
from dataclasses import dataclass, field


def _flatten_table(rows: list[list[str]]) -> str:
    return "\n".join(" | ".join(str(c).strip() for c in row) for row in rows)


@dataclass
class ContractPage:
    page_num: int
    text: str
    tables: list[list[list[str]]] = field(default_factory=list)


@dataclass
class ContractDocument:
    file_path: str
    total_pages: int
    pages: list[ContractPage]
    full_text: str
    full_text_bilingual: str = ""


def extract_contract(file_path: str, keep_english: bool = False) -> ContractDocument:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        return _extract_docx(file_path, keep_english)
    return _extract_pdf(file_path, keep_english)


def _extract_docx(file_path: str, keep_english: bool = False) -> ContractDocument:
    import docx
    doc = docx.Document(file_path)
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    full_text = "\n".join(paragraphs)
    bilingual_text = full_text

    if not keep_english:
        full_text = _strip_english(full_text)

    page = ContractPage(page_num=1, text=full_text)
    return ContractDocument(
        file_path=file_path, total_pages=1, pages=[page],
        full_text=full_text, full_text_bilingual=bilingual_text,
    )


def _extract_pdf(file_path: str, keep_english: bool = False) -> ContractDocument:
    pages: list[ContractPage] = []
    full_text_parts: list[str] = []
    bilingual_parts: list[str] = []

    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            tables = page.extract_tables() or []
            real_tables = [t for t in tables if t and len(t[0]) >= 2]

            pages.append(ContractPage(page_num=i + 1, text=text, tables=real_tables))

            if text:
                bilingual_parts.append(text)
                full_text_parts.append(_strip_english(text) if not keep_english else text)

    full_text = "\n\n".join(full_text_parts)
    full_text_bilingual = "\n\n".join(bilingual_parts)

    return ContractDocument(
        file_path=file_path, total_pages=len(pages), pages=pages,
        full_text=full_text, full_text_bilingual=full_text_bilingual,
    )


def _strip_english(text: str) -> str:
    lines = text.split("\n")
    chinese_lines: list[str] = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            chinese_lines.append("")
            continue
        cjk_count = sum(1 for c in stripped if _is_cjk(c) or c in "，。、；：「」『』（）《》【】")
        ratio = cjk_count / max(len(stripped), 1)
        if ratio >= 0.15:
            chinese_lines.append(stripped)

    return "\n".join(chinese_lines)


def _is_cjk(c: str) -> bool:
    cp = ord(c)
    return (
        (0x4E00 <= cp <= 0x9FFF)
        or (0x3400 <= cp <= 0x4DBF)
        or (0x20000 <= cp <= 0x2A6DF)
    )


def estimate_tokens(text: str) -> int:
    cjk = sum(1 for c in text if _is_cjk(c))
    other = len(text) - cjk
    return int(cjk * 1.0 + other * 0.75)


def extract_table_text(tables: list[list[list[str]]]) -> str:
    parts: list[str] = []
    for table in tables:
        parts.append("--- TABLE ---")
        for row in table:
            cells = [str(c).strip() if c else "" for c in row]
            parts.append(" | ".join(cells))
        parts.append("--- END TABLE ---")
    return "\n".join(parts)
